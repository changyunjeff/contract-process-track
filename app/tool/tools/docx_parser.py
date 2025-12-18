"""
DOCX 文件解析工具
用于将 DOCX 文件解析为纯文本
"""
from typing import Dict, Any
from .base import ToolInterface

try:
    from docx import Document
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False


class DocxParserTool(ToolInterface):
    name = "docx_parser"
    description = "Parse DOCX files and extract text content. Supports extracting text from Word documents (.docx format)."

    def get_parameters_schema(self) -> Dict[str, Any]:
        """获取DOCX解析工具的JSON Schema定义"""
        return {
            "type": "object",
            "properties": {
                "file_path": {
                    "type": "string",
                    "description": "Path to the DOCX file to be parsed"
                },
                "include_headers": {
                    "type": "boolean",
                    "description": "Whether to include headers and footers in the extracted text (default: false)",
                    "default": False
                },
                "include_tables": {
                    "type": "boolean",
                    "description": "Whether to include table content in the extracted text (default: true)",
                    "default": True
                }
            },
            "required": ["file_path"]
        }

    def run(self, input: dict, context: dict) -> dict:
        """
        执行DOCX文件解析
        
        Args:
            input: 包含file_path和可选参数的字典
            context: 执行上下文
            
        Returns:
            dict: 包含提取文本的结果
        """
        if not DOCX_AVAILABLE:
            return {
                "type": "error",
                "error": "python-docx library is not installed. Please install it using: pip install python-docx"
            }
        
        file_path = input.get("file_path")
        if not file_path:
            raise ValueError("file_path is required")
        
        include_headers = input.get("include_headers", False)
        include_tables = input.get("include_tables", True)
        
        # 如果路径不是绝对路径，则使用基础路径（与 OCR 工具保持一致）
        import os
        if not os.path.isabs(file_path):
            base_path = "/data/files"
            full_path = os.path.join(base_path, file_path)
        else:
            full_path = file_path
        
        try:
            # 打开 DOCX 文件
            doc = Document(full_path)
            
            # 提取段落文本
            paragraphs = []
            for paragraph in doc.paragraphs:
                text = paragraph.text.strip()
                if text:
                    paragraphs.append(text)
            
            # 提取表格内容（如果需要）
            tables_text = []
            if include_tables:
                for table in doc.tables:
                    table_rows = []
                    for row in table.rows:
                        row_cells = []
                        for cell in row.cells:
                            cell_text = cell.text.strip()
                            if cell_text:
                                row_cells.append(cell_text)
                        if row_cells:
                            table_rows.append(" | ".join(row_cells))
                    if table_rows:
                        tables_text.append("\n".join(table_rows))
            
            # 提取页眉页脚（如果需要）
            headers_footers_text = []
            if include_headers:
                for section in doc.sections:
                    # 页眉
                    if section.header:
                        for paragraph in section.header.paragraphs:
                            text = paragraph.text.strip()
                            if text:
                                headers_footers_text.append(f"[Header] {text}")
                    # 页脚
                    if section.footer:
                        for paragraph in section.footer.paragraphs:
                            text = paragraph.text.strip()
                            if text:
                                headers_footers_text.append(f"[Footer] {text}")
            
            # 组合所有文本
            all_text_parts = []
            
            if headers_footers_text:
                all_text_parts.extend(headers_footers_text)
                all_text_parts.append("")  # 添加分隔空行
            
            all_text_parts.extend(paragraphs)
            
            if tables_text:
                all_text_parts.append("")  # 添加分隔空行
                all_text_parts.append("[Tables]")
                all_text_parts.extend(tables_text)
            
            full_text = "\n".join(all_text_parts)
            
            # 统计信息
            stats = {
                "paragraph_count": len(paragraphs),
                "table_count": len(doc.tables) if include_tables else 0,
                "total_characters": len(full_text),
                "total_words": len(full_text.split()) if full_text else 0
            }
            
            return {
                "type": "text",
                "source": file_path,
                "content": full_text,
                "stats": stats,
                "include_headers": include_headers,
                "include_tables": include_tables
            }
        except FileNotFoundError:
            return {
                "type": "error",
                "source": file_path,
                "error": f"File not found: {full_path}"
            }
        except Exception as e:
            return {
                "type": "error",
                "source": file_path,
                "error": f"Failed to parse DOCX file: {str(e)}"
            }

